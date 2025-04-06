import sqlite3
import streamlit as st
import pandas as pd
import google.generativeai as genai
import matplotlib.pyplot as plt
import io
import os
import json
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import requests
from googleapiclient.discovery import build
from bs4 import BeautifulSoup
import time
import re
import hashlib
import uuid
import base64
from PIL import Image
from pathlib import Path

# API Keys - Replace with your own keys
GOOGLE_API_KEY = "AIzaSyCDVt_gW-DFZcHmxCBUkQnLXtw6uXXaCe8"
GOOGLE_CSE_ID = "a59d66d6326034ab6"
GEMINI_API_KEY = "AIzaSyBmkHHRyGXbC6_TM5wI_p9qrhIC4LJKzzU"

# Initialize Gemini model
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-pro')

# ------- Database Functions -------

def init_chat_database(username):
    """Initialize or connect to the user's chat database"""
    user_dir = get_user_data_dir(username)
    db_path = user_dir / "chat_history.db"
    
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    # Create messages table if it doesn't exist
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS messages (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        role TEXT NOT NULL,
        content TEXT NOT NULL,
        timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    
    conn.commit()
    return conn

def save_chat_history_to_db(username):
    """Save current chat history to SQLite database"""
    if not is_authenticated() or not st.session_state.messages:
        return False
    
    try:
        # Connect to database
        conn = init_chat_database(username)
        cursor = conn.cursor()
        
        # Clear existing messages to avoid duplicates
        cursor.execute("DELETE FROM messages")
        
        # Insert all current messages
        for msg in st.session_state.messages:
            cursor.execute(
                "INSERT INTO messages (role, content) VALUES (?, ?)",
                (
                    msg["role"],
                    msg.get("content", "")
                )
            )
        
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Error saving chat history to database: {e}")
        return False

def load_chat_history_from_db(username):
    """Load chat history from SQLite database"""
    if not username:
        return []
    
    try:
        user_dir = get_user_data_dir(username)
        db_path = user_dir / "chat_history.db"
        
        if not db_path.exists():
            return []
            
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Fetch all messages ordered by id (chronological order)
        cursor.execute("SELECT role, content FROM messages ORDER BY id")
        messages = []
        
        for row in cursor.fetchall():
            role, content = row
            message = {"role": role, "content": content}
            messages.append(message)
        
        conn.close()
        return messages
    except Exception as e:
        st.error(f"Error loading chat history from database: {e}")
        return []

def clear_chat_history_db(username):
    """Clear all chat history for a user from the database"""
    if not username:
        return False
    
    try:
        user_dir = get_user_data_dir(username)
        db_path = user_dir / "chat_history.db"
        
        if not db_path.exists():
            return True
            
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Delete all messages
        cursor.execute("DELETE FROM messages")
        
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Error clearing chat history from database: {e}")
        return False

# ------- Authentication System -------

def authenticate(username, password):
    """Authenticate a user and load their previous chat history"""
    users = load_users()
    if username in users and users[username]["password"] == hash_password(password):
        st.session_state.authenticated = True
        st.session_state.username = username
        
        # Load previous chat history from database
        if "messages" in st.session_state:
            prev_messages = load_chat_history_from_db(username)
            if prev_messages:
                st.session_state.messages = prev_messages
                st.session_state.db_loaded = True  # Flag to indicate DB was loaded
            else:
                st.session_state.messages = []
                st.session_state.db_loaded = False
        
        return True
    return False

def load_users():
    """Load existing users from the JSON file"""
    if os.path.exists("users.json"):
        with open("users.json", "r") as f:
            return json.load(f)
    return {}

def save_users(users):
    """Save users to the JSON file"""
    with open("users.json", "w") as f:
        json.dump(users, f)

def hash_password(password):
    """Hash password for secure storage"""
    return hashlib.sha256(password.encode()).hexdigest()

def is_authenticated():
    """Check if user is authenticated"""
    return 'authenticated' in st.session_state and st.session_state.authenticated

def create_account(username, password, email):
    """Create a new user account"""
    users = load_users()
    if username in users:
        return False, "Username already exists"
    
    # Validate email format (basic check)
    if "@" not in email or "." not in email:
        return False, "Invalid email format"
    
    # Validate password strength
    if len(password) < 6:
        return False, "Password must be at least 6 characters"
    
    # Create user
    users[username] = {
        "password": hash_password(password),
        "email": email,
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
    }
    save_users(users)
    return True, "Account created successfully!"

def logout():
    """Log out the current user"""
    for key in list(st.session_state.keys()):
        del st.session_state[key]

# ------- UI Helper Functions -------

def add_logo_and_background():
    """Add logo and background to the page"""
    st.markdown(
        """
        <style>
        .main {
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        }
        .auth-container {
            background-color: white;
            padding: 2rem;
            border-radius: 10px;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
            max-width: 450px;
            margin: 0 auto;
            margin-top: 2rem;
        }
        .auth-header {
            text-align: center;
            margin-bottom: 2rem;
        }
        .auth-form {
            margin-bottom: 1rem;
        }
        .auth-button {
            width: 100%;
            margin-top: 1rem;
        }
        .toggle-link {
            text-align: center;
            margin-top: 1rem;
            cursor: pointer;
            color: #4e6bba;
        }
        .toggle-link:hover {
            text-decoration: underline;
        }
        .app-title {
            font-family: 'Arial', sans-serif;
            font-weight: bold;
            font-size: 2.5rem;
            color: #2c3e50;
            text-align: center;
            margin-bottom: 0.5rem;
        }
        .app-subtitle {
            font-family: 'Arial', sans-serif;
            font-size: 1.2rem;
            color: #7f8c8d;
            text-align: center;
            margin-bottom: 2rem;
        }
        .stApp > header {
            background-color: transparent;
        }
        .stTextInput > div > div > input {
            padding: 12px 15px;
            border-radius: 8px;
        }
        .stButton > button {
            border-radius: 8px;
            padding: 10px 15px;
            font-weight: 600;
            background-color: #4361ee;
            color: white;
            border: none;
            transition: all 0.3s;
        }
        .stButton > button:hover {
            background-color: #3a56d4;
            box-shadow: 0 4px 10px rgba(0, 0, 0, 0.1);
        }
        </style>
        """,
        unsafe_allow_html=True
    )

# ------- Login Page Function -------

def show_login_page():
    """Display the login page"""
    add_logo_and_background()
    
    st.markdown("<h1 class='app-title'>AI Document Assistant</h1>", unsafe_allow_html=True)
    st.markdown("<p class='app-subtitle'>Your AI-powered knowledge companion</p>", unsafe_allow_html=True)
    
    # Login container
    st.markdown("<div class='auth-container'>", unsafe_allow_html=True)
    st.markdown("<h2 class='auth-header'>Login</h2>", unsafe_allow_html=True)
    
    username = st.text_input("Username", key="login_username")
    password = st.text_input("Password", type="password", key="login_password")
    
    col1, col2 = st.columns([3, 1])
    with col1:
        login_button = st.button("Login", use_container_width=True)
    with col2:
        if st.button("Demo", use_container_width=True):
            # Set demo mode - instant access with limited features
            st.session_state.authenticated = True
            st.session_state.username = "demo_user"
            st.session_state.demo_mode = True
            st.rerun()
    
    if login_button:
        if authenticate(username, password):
            st.session_state.demo_mode = False
            st.rerun()
        else:
            st.error("Invalid username or password")
    
    st.markdown("<div class='toggle-link'>", unsafe_allow_html=True)
    if st.button("Don't have an account? Sign up"):
        st.session_state.show_signup = True
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Add some information about the app below the login form
    st.markdown("---")
    st.markdown("### About AI Document Assistant")
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
        - Ask any question in natural language
        - Upload and analyze DOCX files
        - Find information from the web
        - Get detailed explanations
        """)
    with col2:
        st.markdown("""
        - Powered by Google's Gemini AI
        - Access knowledge from multiple sources
        - Save and export your conversations
        - Rate answers to improve quality
        """)

# ------- Signup Page Function -------

def show_signup_page():
    """Display the signup page"""
    add_logo_and_background()
    
    st.markdown("<h1 class='app-title'>AI Document Assistant</h1>", unsafe_allow_html=True)
    st.markdown("<p class='app-subtitle'>Your AI-powered knowledge companion</p>", unsafe_allow_html=True)
    
    # Signup container
    st.markdown("<div class='auth-container'>", unsafe_allow_html=True)
    st.markdown("<h2 class='auth-header'>Create an Account</h2>", unsafe_allow_html=True)
    
    username = st.text_input("Username", key="signup_username")
    email = st.text_input("Email Address")
    password = st.text_input("Password", type="password", key="signup_password")
    confirm_password = st.text_input("Confirm Password", type="password")
    
    signup_button = st.button("Sign Up", use_container_width=True)
    
    if signup_button:
        if password != confirm_password:
            st.error("Passwords do not match")
        else:
            success, message = create_account(username, password, email)
            if success:
                st.success(message)
                # Auto-login after successful signup
                authenticate(username, password)
                st.session_state.demo_mode = False
                st.rerun()
            else:
                st.error(message)
    
    st.markdown("<div class='toggle-link'>", unsafe_allow_html=True)
    if st.button("Already have an account? Login"):
        st.session_state.show_signup = False
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)
    
    st.markdown("</div>", unsafe_allow_html=True)

# ------- Profile Component -------

def show_profile_component():
    """Display user profile in the sidebar"""
    if is_authenticated():
        users = load_users()
        user_data = users.get(st.session_state.username, {})
        
        with st.sidebar:
            st.sidebar.title("User Profile")
            st.sidebar.write(f"Username: {st.session_state.username}")
            
            if 'email' in user_data:
                st.sidebar.write(f"Email: {user_data['email']}")
            
            if 'created_at' in user_data:
                st.sidebar.write(f"Account created: {user_data['created_at']}")
            
            if st.sidebar.button("Logout"):
                logout()
                st.rerun()
            
            # Add some visual separator
            st.sidebar.markdown("---")

# ------- User-Specific Data Management -------

def get_user_data_dir(username):
    """Create and return user-specific data directory"""
    user_dir = Path(f"user_data/{username}")
    user_dir.mkdir(parents=True, exist_ok=True)
    return user_dir

def save_chat_history_to_docx(username):
    """Save the current chat history to a DOCX file for the user"""
    if not is_authenticated() or not st.session_state.messages:
        return False
    
    try:
        # Create user directory
        user_dir = get_user_data_dir(username)
        
        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading(f'Chat History for {username}', 0)
        doc.add_paragraph(f'Generated on: {time.strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')
        
        # Add each message
        for msg in st.session_state.messages:
            # Add role as a heading
            doc.add_heading(f"{msg['role'].capitalize()}", level=2)
            
            # Add content
            doc.add_paragraph(msg['content'])
            
            # Add a separator
            doc.add_paragraph('---')
        
        # Save the document
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        filename = f"chat_history_{timestamp}.docx"
        filepath = user_dir / filename
        doc.save(filepath)
        
        return str(filepath)
    except Exception as e:
        st.error(f"Error saving chat history: {e}")
        return False

def get_user_embedding_dir(username):
    """Get the embedding directory path for a specific user"""
    user_dir = get_user_data_dir(username)
    embed_dir = user_dir / "embeddings"
    embed_dir.mkdir(exist_ok=True)
    return embed_dir

# ------- Document Processing -------

def process_uploaded_documents(uploaded_files):
    """Process uploaded DOCX files and create embeddings"""
    if not uploaded_files:
        return False
    
    all_text = ""
    file_summaries = []
    total_paragraphs = 0
    
    for uploaded_file in uploaded_files:
        try:
            # Save the uploaded file temporarily
            file_path = f"temp_{uploaded_file.name}"
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Process the DOCX file
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Add document metadata
            file_text = f"DOCUMENT: {uploaded_file.name}\n{text}\nEND OF DOCUMENT: {uploaded_file.name}\n\n"
            all_text += file_text
            
            # Create summary for this file
            file_summaries.append({
                'filename': uploaded_file.name,
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            })
            
            total_paragraphs += len(doc.paragraphs)
            
            # Clean up temp file
            os.remove(file_path)
            
        except Exception as e:
            st.error(f"Error processing {uploaded_file.name}: {str(e)}")
    
    # If no files were successfully processed
    if not file_summaries:
        st.error("No DOCX files could be processed.")
        return False
    
    # Split all text into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(all_text)
    
    # Store in session state
    st.session_state.doc_data = {
        'text': all_text,
        'chunks': chunks,
        'files': file_summaries,
        'summary': f"Processed {len(file_summaries)} documents with {total_paragraphs} paragraphs and {len(chunks)} text chunks."
    }
    
    # Create embeddings based on user status
    if is_authenticated() and not st.session_state.demo_mode:
        username = st.session_state.username
        user_embed_dir = get_user_embedding_dir(username)
        
        try:
            with st.spinner("Creating embeddings for your documents..."):
                embeddings_model = get_embeddings_model(username)
                st.session_state.vectorstore = Chroma(
                    persist_directory=str(user_embed_dir),
                    embedding_function=embeddings_model
                )
                # Add documents to vectorstore
                st.session_state.vectorstore.add_texts(
                    texts=chunks
                )
                st.success("Document embeddings created successfully!")
        except Exception as e:
            st.error(f"Failed to create document embeddings: {e}")
            return False
    else:
        # For demo mode, create in-memory embeddings
        try:
            with st.spinner("Creating document embeddings..."):
                embeddings_model = get_embeddings_model()
                st.session_state.vectorstore = Chroma(
                    embedding_function=embeddings_model
                )
                # Add documents to vectorstore
                st.session_state.vectorstore.add_texts(
                    texts=chunks
                )
                st.success("Document embeddings created successfully!")
        except Exception as e:
            st.error(f"Failed to create document embeddings: {e}")
            return False
    
    return True

@st.cache_resource
def get_embeddings_model(username=None):
    """Initialize and cache the embedding model with user-specific caching"""
    try:
        # Use the same model but note that we're using different parameters
        # to ensure a unique cache entry per user
        model = HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'},  # Changed to CPU for compatibility
            encode_kwargs={'normalize_embeddings': True}
        )
        
        # Test the model
        test_embedding = model.embed_query("test")
        if test_embedding is None:
            raise ValueError("Embeddings model failed to generate embeddings")
        return model
    except Exception as e:
        st.error(f"Failed to initialize embeddings model: {str(e)}")
        return None

def get_web_content(query, num_results=3):
    """Get relevant content using Google Custom Search API"""
    web_content = []
    
    try:
        # Initialize the Custom Search API service
        service = build(
            "customsearch", "v1",
            developerKey=GOOGLE_API_KEY
        )
        
        # Perform the search
        result = service.cse().list(
            q=query,
            cx=GOOGLE_CSE_ID,
            num=num_results
        ).execute()
        
        # Process search results
        if 'items' in result:
            for item in result['items']:
                web_content.append({
                    'url': item['link'],
                    'title': item.get('title', ''),
                    'content': item.get('snippet', ''),
                    'source': 'Web Search'
                })
                
        return web_content
    except Exception as e:
        st.warning(f"Web search failed: {str(e)}")
        return []

def get_relevant_chunks(query, doc_data):
    """Retrieve relevant text chunks using embeddings and similarity search."""
    if not doc_data or 'chunks' not in doc_data or not doc_data['chunks']:
        return []

    try:
        username = st.session_state.username if is_authenticated() else None
        
        # Perform similarity search
        if "vectorstore" in st.session_state and st.session_state.vectorstore is not None:
            try:
                results = st.session_state.vectorstore.similarity_search(
                    query,
                    k=5  # Number of results to retrieve
                )
                if not results:
                    return []
                
                return [doc.page_content for doc in results]
            except Exception as e:
                st.warning(f"Document search failed: {str(e)}")
                return []
        else:
            return []
    except Exception as e:
        st.warning(f"Error processing document search: {str(e)}")
        return []

def generate_response(query, doc_chunks=None, web_content=None):
    """Generate response using Gemini model with document and web context if available"""
    
    # Capture recent conversation history
    chat_history = ""
    if len(st.session_state.messages) > 0:
        # Get the last 5 messages or all messages if less than 5
        recent_messages = st.session_state.messages[-5:] if len(st.session_state.messages) > 5 else st.session_state.messages
        chat_history = "\n".join([f"{msg['role'].capitalize()}: {msg['content']}" for msg in recent_messages])
    
    # Base prompt for general questions
    if not doc_chunks and not web_content:
        prompt = f"""
You are a helpful, friendly AI assistant capable of answering questions on a wide range of topics.

Recent conversation history:
{chat_history}

User Question: {query}

If the question refers to previous messages or conversation history, please reference the conversation history above.

Please provide a detailed, informative response. Use markdown formatting to structure your answer clearly.
"""
    else:
        # Create document context if available
        doc_context = ""
        if doc_chunks and len(doc_chunks) > 0:
            doc_context = "Document Content:\n" + "\n\n".join(doc_chunks)
        
        # Create web context if available
        web_context = ""
        if web_content and len(web_content) > 0:
            web_sources = [f"Source: {content['url']}\nTitle: {content['title']}\nContent: {content['content']}" for content in web_content]
            web_context = "Web Sources:\n" + "\n\n".join(web_sources)
        
        # Enhanced prompt with document and web content
        prompt = f"""
You are a helpful, friendly AI assistant with access to document content and web search results.

Recent conversation history:
{chat_history}

User Question: {query}

{doc_context}

{web_context}

If the query is about previous messages or conversation history, please refer to the conversation history above.

Please provide a detailed, comprehensive response that:
1. Directly answers the user's question with the most relevant information
2. Incorporates insights from the documents and web sources when relevant
3. Uses clear, engaging language and well-structured markdown formatting 
4. Cites sources when using specific information from documents or web
5. Maintains a conversational, helpful tone

If sources conflict, acknowledge this and provide the most reliable information while noting the discrepancy.
"""

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error generating response: {e}")
        return f"I'm sorry, I encountered an error while generating a response. Please try again or rephrase your question. Error details: {str(e)}"

def display_chat_message(role, content, message_id=None):
    """Display a chat message with optional star rating."""
    with st.chat_message(role):
        st.markdown(content)
        
        # Add star rating if this is an assistant message
        if role == "assistant" and message_id is not None:
            # Create a unique key for this message's feedback
            feedback_key = f"feedback_{message_id}"
            
            # Check if feedback was already given for this message
            if str(message_id) in st.session_state.get("feedback_data", {}) and "rating" in st.session_state.feedback_data[str(message_id)]:
                rating = st.session_state.feedback_data[str(message_id)]["rating"]
                st.success(f"You rated this response: {rating}/5 stars")
            else:
                st.markdown("**How helpful was this response?**")
                col1, col2, col3, col4, col5 = st.columns(5)
                
                if col1.button("⭐ 1", key=f"star_1_{message_id}"):
                    rate_response(str(message_id), 1)
                    st.rerun()
                    
                if col2.button("⭐ 2", key=f"star_2_{message_id}"):
                    rate_response(str(message_id), 2)
                    st.rerun()
                    
                if col3.button("⭐ 3", key=f"star_3_{message_id}"):
                    rate_response(str(message_id), 3)
                    st.rerun()
                    
                if col4.button("⭐ 4", key=f"star_4_{message_id}"):
                    rate_response(str(message_id), 4)
                    st.rerun()
                    
                if col5.button("⭐ 5", key=f"star_5_{message_id}"):
                    rate_response(str(message_id), 5)
                    st.rerun()

def rate_response(message_id, star_num):
    """Handle star rating logic"""
    # Get the current query
    current_query = st.session_state.get("current_query", "Unknown query")
    
    # Initialize feedback data if not present
    if "feedback_data" not in st.session_state:
        st.session_state.feedback_data = {}
    
    # Save the rating
    if message_id not in st.session_state.feedback_data:
        st.session_state.feedback_data[message_id] = {}
        
    # Get the message data (convert message_id to integer for indexing)
    message = st.session_state.messages[int(message_id)]
    
    # Store feedback data
    st.session_state.feedback_data[message_id] = {
        "query": current_query,
        "content": message.get("content", ""),
        "rating": star_num,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
    }
    
    # If we want to save feedback to file
    if is_authenticated() and not st.session_state.demo_mode:
        try:
            # User-specific feedback file
            username = st.session_state.username
            user_dir = get_user_data_dir(username)
            feedback_file = user_dir / "feedback_data.json"
            
            # Load existing feedback if available
            existing_feedback = {}
            if feedback_file.exists():
                with open(feedback_file, 'r') as f:
                    existing_feedback = json.load(f)
                    
            # Update with new feedback
            existing_feedback.update(st.session_state.feedback_data)
            
            # Save updated feedback
            with open(feedback_file, 'w') as f:
                json.dump(existing_feedback, f)
                
        except Exception as e:
            st.error(f"Error saving feedback: {e}")

def main():
    # Configure page settings
    st.set_page_config(page_title="AI Document Assistant", layout="wide")
    
    # Initialize session state variables for authentication
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    
    if "demo_mode" not in st.session_state:
        st.session_state.demo_mode = False
        
    if "show_signup" not in st.session_state:
        st.session_state.show_signup = False
        
    if "last_autosave" not in st.session_state:
        st.session_state.last_autosave = time.time()
    
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    if "db_loaded" not in st.session_state:
        st.session_state.db_loaded = False
        
    if "doc_data" not in st.session_state:
        st.session_state.doc_data = None
        
    if "vectorstore" not in st.session_state:
        st.session_state.vectorstore = None
        
    if "feedback_data" not in st.session_state:
        st.session_state.feedback_data = {}
        
    if "current_query" not in st.session_state:
        st.session_state.current_query = None
    
    # Display the appropriate page based on authentication state
    if not is_authenticated():
        if st.session_state.show_signup:
            show_signup_page()
        else:
            show_login_page()
    else:
        # If authenticated, show the main app
        # Show profile in sidebar
        show_profile_component()
        
        # Display demo mode warning if applicable
        if st.session_state.demo_mode:
            st.warning("You are in demo mode. Some features may be limited. Sign up for full access.")
        elif st.session_state.db_loaded:
            st.success("Previous chat history loaded successfully.")
            # Reset the flag after showing the message
            st.session_state.db_loaded = False
        
        # Main app title
        st.title("AI Document Assistant")
        st.caption(f"Welcome, {st.session_state.username}! Ask me anything or upload documents for more context.")
        
        # Sidebar for document upload and management
        with st.sidebar:
            st.markdown("### Upload Documents")
            uploaded_files = st.file_uploader(
                "Upload DOCX files", 
                type=["docx"], 
                accept_multiple_files=True
            )
            
            if uploaded_files:
                if st.button("Process Documents"):
                    success = process_uploaded_documents(uploaded_files)
                    if success:
                        st.success("Documents processed successfully!")
                    
            if st.session_state.doc_data:
                st.markdown("### Processed Documents")
                for idx, file_info in enumerate(st.session_state.doc_data["files"]):
                    st.write(f"{idx+1}. {file_info['filename']}")
                
                if st.button("Clear Documents"):
                    st.session_state.doc_data = None
                    st.session_state.vectorstore = None
                    st.success("Documents cleared.")
                    st.rerun()
            
            # Add export options for non-demo users
            if not st.session_state.demo_mode:
                st.markdown("### Export Data")
                if st.button("Export Chat History"):
                    file_path = save_chat_history_to_docx(st.session_state.username)
                    if file_path:
                        st.success(f"Chat history saved to: {file_path}")
            
            # Chat controls
            st.markdown("### Chat Controls")
            if st.button("Clear Chat History"):
                st.session_state.messages = []
                if is_authenticated() and not st.session_state.demo_mode:
                    clear_chat_history_db(st.session_state.username)
                st.success("Chat history cleared.")
                st.rerun()
        
        # Main chat area
        # Display chat history
        for i, message in enumerate(st.session_state.messages):
            display_chat_message(
                message["role"],
                message["content"],
                message_id=i if message["role"] == "assistant" else None
            )
        
        # Search options area
        search_option = "both"
        if st.session_state.doc_data:
            search_cols = st.columns([1, 1, 1])
            with search_cols[0]:
                doc_search = st.checkbox("Search Documents", value=True)
            with search_cols[1]:
                web_search = st.checkbox("Search Web", value=True)
            with search_cols[2]:
                gemini_ai = st.checkbox("Use Gemini AI", value=True)
                
            if doc_search and web_search:
                search_option = "both"
            elif doc_search:
                search_option = "documents"
            elif web_search:
                search_option = "web"
            else:
                search_option = "none"
                if gemini_ai:
                    search_option = "ai_only"
        
        # Chat input
        query = st.chat_input("Ask me anything...")
        
        if query:
            st.session_state.current_query = query
            st.session_state.messages.append({"role": "user", "content": query})
            
            with st.spinner("Thinking..."):
                # Search documents if available
                doc_chunks = []
                if st.session_state.doc_data and (search_option == "documents" or search_option == "both"):
                    doc_chunks = get_relevant_chunks(query, st.session_state.doc_data)
                
                # Search web if enabled
                web_content = []
                if search_option == "web" or search_option == "both":
                    web_content = get_web_content(query)
                
                # Generate response
                response = generate_response(query, doc_chunks, web_content)
                
                if response:
                    message_id = len(st.session_state.messages)
                    st.session_state.messages.append({
                        "role": "assistant",
                        "content": response
                    })
                    
                    # Auto-save after new message for registered users
                    if is_authenticated() and not st.session_state.demo_mode:
                        save_chat_history_to_db(st.session_state.username)
                    
                    # Rerun to display the new message
                    st.rerun()

if __name__ == "__main__":
    main()